from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_type, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{margin_type}"))
        if node is None:
            node = OxmlElement(f"w:{margin_type}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Calibri") -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 15, "000000", 12, 12),
        ("Heading 2", 14, "000000", 6, 6),
        ("Heading 3", 12, "000000", 0, 0),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0 if name in ("Heading 1", "Heading 2") else 1.25


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Menlo"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        else:
            run.text = part
            set_run_font(run)


def add_code_block(doc: Document, code: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.18)
    para.paragraph_format.right_indent = Inches(0.18)
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(5)
    run = para.add_run(code)
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F8FAFC")
    para._p.get_or_add_pPr().append(shading)


def add_image(doc: Document, rel_path: str) -> None:
    image_path = ROOT / rel_path
    if not image_path.exists():
        return
    with Image.open(image_path) as img:
        width, height = img.size
    max_width = Inches(6.5)
    max_height = Inches(3.2)
    ratio = min(max_width / width, max_height / height)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=int(width * ratio), height=int(height * ratio))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    set_cell_margins(table)
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            add_inline_markdown(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(10.5)
                set_run_font(run)
                if row_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if row_idx == 0:
                set_cell_shading(cell, "1F2937")


def parse_table(lines: list[str], index: int) -> tuple[list[list[str]], int]:
    table_lines = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows, index


def markdown_to_docx(markdown: str, out_path: Path) -> None:
    doc = Document()
    set_doc_defaults(doc)
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        image = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image:
            add_image(doc, image.group(2))
            index += 1
            continue

        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(style=None)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor.from_string("0F172A")
            set_run_font(run, east_asia="黑体")
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:])
        elif re.match(r"\d+\. ", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, re.sub(r"^\d+\. ", "", stripped))
        else:
            paragraph = doc.add_paragraph()
            add_inline_markdown(paragraph, stripped)
        index += 1

    doc.save(out_path)


def main() -> None:
    markdown = (ROOT / "report.md").read_text(encoding="utf-8")
    markdown_to_docx(markdown, ROOT / "report.docx")


if __name__ == "__main__":
    main()
